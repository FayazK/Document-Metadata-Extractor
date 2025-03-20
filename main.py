import os
import logging
import docx
import PyPDF2
import google.generativeai as genai
import openpyxl
import xml.etree.ElementTree as ET
from pathlib import Path
import sys
import argparse
from typing import Dict, Any, Optional, Tuple
import json

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("metadata_extraction.log"),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger("metadata_extractor")


# Function to read the API key from a config file
def load_api_key(config_path: str = "config.json") -> str:
    """Load the API key from a configuration file."""
    try:
        with open(config_path, 'r') as f:
            config = json.load(f)
            return config.get("api_key", "")
    except (FileNotFoundError, json.JSONDecodeError) as e:
        logger.error(f"Error loading API key: {e}")
        return ""


class DocumentParser:
    """Base class for document parsing functionality."""

    @staticmethod
    def get_file_size(file_path: str) -> str:
        """Get the file size in KB or MB."""
        size_bytes = os.path.getsize(file_path)
        if size_bytes < 1024 * 1024:  # Less than 1MB
            return f"{size_bytes / 1024:.2f} KB"
        else:
            return f"{size_bytes / (1024 * 1024):.2f} MB"

    @staticmethod
    def extract_text_from_docx(file_path: str) -> Optional[str]:
        """Extract text content from a DOCX file."""
        try:
            doc = docx.Document(file_path)
            full_text = []

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():  # Only add non-empty paragraphs
                    full_text.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        full_text.append(" | ".join(row_text))

            return '\n'.join(full_text)
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {e}")
            return None

    @staticmethod
    def extract_text_from_pdf(file_path: str) -> Optional[str]:
        """Extract text content from a PDF file."""
        try:
            text = []
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                num_pages = len(reader.pages)

                logger.info(f"Extracting text from PDF with {num_pages} pages")
                for i, page in enumerate(reader.pages):
                    page_text = page.extract_text() or ""
                    if page_text.strip():
                        text.append(page_text)

                    # Log progress for large PDFs
                    if (i + 1) % 10 == 0:
                        logger.info(f"Processed {i + 1}/{num_pages} pages")

            return '\n'.join(text)
        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {e}")
            return None


class MetadataExtractor:
    """Class responsible for extracting metadata using Gemini API."""

    def __init__(self, api_key: str, output_dir: str = "output"):
        """Initialize the metadata extractor with the API key and output directory."""
        self.api_key = api_key
        self.output_dir = output_dir
        self.configure_api()
        self.prompt = self._create_prompt_template()

        # Create output directory if it doesn't exist
        os.makedirs(self.output_dir, exist_ok=True)

    def configure_api(self) -> None:
        """Configure the Gemini API with the provided key."""
        try:
            genai.configure(api_key=self.api_key)
            self.model = genai.GenerativeModel('gemini-2.0-flash')
            logger.info("Gemini API configured successfully")
        except Exception as e:
            logger.error(f"Error configuring Gemini API: {e}")
            raise RuntimeError(f"Failed to configure Gemini API: {e}")

    def _create_prompt_template(self) -> str:
        """Create the prompt template for metadata extraction."""
        return """You are an expert metadata extraction system. Your task is to analyze the provided document text and extract relevant metadata based on the following schema. You must return the extracted metadata in XML format.

**Metadata Schema:**

**I. General Metadata (Applicable to Most Document Types):**
* **DocumentTitle:** The official title of the document.
* **DocumentType:** (Choose one from: Research Paper, Test Conducted, Test Result, FAQ, EPD, ASTM Standard, Aramid Data Sheet, BATT Test Data, Case Study, Lifecycle Cost Analysis, Technical Product Data)
* **DatePublishedCreated:** The date the document was officially published or created (YYYY-MM-DD format if possible).
* **VersionRevisionNumber:** If applicable, the version or revision number of the document.
* **SourceOrigin:** Where the document originated (e.g., specific lab, department, external organization, ASTM).
* **Keywords:** Relevant keywords or tags describing the document's content separated by comma.
* **AbstractSummary:** A brief summary of the document's main points or findings.
* **FileFormat:** (e.g., PDF, DOCX)
* **FileSizeKBMB:** The size of the document file in KB or MB.
* **DocumentIDReferenceNumber:** A unique identifier assigned to the document (if applicable).
* **ConfidentialityLevel:** (e.g., Public, Internal, Confidential).
* **RelevantProductsMaterials:** Specific Surface Tech asphalt products or materials discussed.
* **GeographicRegion:** If the document specifically pertains to a certain geographic location or project.

**II. Specific Metadata (Tailored to Each Document Type):**

* **Research Paper:**
    * **Authors:** List of authors involved.
    * **JournalConferenceName:** If published in a journal or presented at a conference.
    * **DOI:** If available.
    * **ResearchMethodology:** Briefly describe the methodology used.
    * **KeyFindings:** Summarize the most significant results.
    * **StudyAreaLocation:** If the research focused on a specific area.

* **Test Conducted:**
    * **TestNameType:** Specific name of the test performed (e.g., "Rutting Test").
    * **TestStandardUsed:** The specific ASTM or other standard followed (e.g., "ASTM D6927").
    * **DateOfTest:** Date when the test was conducted (YYYY-MM-DD format if possible).
    * **TestingLaboratory:** Name of the lab where the test was performed.
    * **EquipmentUsed:** List of key equipment used.
    * **MaterialsTested:** Detailed description of the materials tested.

* **Test Result:**
    * **TestNameTypeReference:** Link or reiterate the test name.
    * **DateOfTest:** Date when the test was conducted (YYYY-MM-DD format if possible).
    * **KeyResultsDataPoints:** Specific numerical results or observations.
    * **PassFailCriteria:** If the test had specific acceptance criteria.
    * **ConclusionsInterpretations:** The main conclusions drawn from the results.

* **FAQ:**
    * **Question:** The specific question asked.
    * **Answer:** The detailed answer provided.
    * **RelatedTopics:** Keywords or categories the FAQ falls under.

* **EPD:**
    * **DeclaredUnit:** The functional unit for which environmental data is declared.
    * **GWP:** The reported Global Warming Potential value.
    * **LCAPractitioner:** Name of the organization or individual who performed the LCA.
    * **PCR:** The specific Product Category Rules document used.
    * **ValidityPeriod:** The start and end dates for which the EPD is valid (YYYY-MM-DD format if possible).

* **ASTM Standard:**
    * **ASTMDesignation:** The unique identifier for the standard (e.g., "ASTM D6927-15").
    * **YearOfIssue:** The year the standard was published or last revised.
    * **TitleOfStandard:** The official title of the ASTM standard.
    * **RelevantSections:** If the document is referencing specific sections relevant to Surface Tech.

* **Aramid Data Sheet:**
    * **AramidTypeGrade:** Specific type or grade of aramid fiber.
    * **KeyProperties:** Important technical properties (e.g., tensile strength, modulus).
    * **Supplier:** The manufacturer or supplier of the aramid material.
    * **ApplicationArea:** Where this specific aramid data is relevant.

* **BATT Test Data:**
    * **TestLocation:** Where the BATT test was conducted.
    * **DateOfTest:** Date of the BATT test (YYYY-MM-DD format if possible).
    * **BATTParametersMeasured:** List the specific parameters measured.
    * **ResultsSummary:** A brief summary of the BATT test results.

* **Case Study:**
    * **ProjectNameLocation:** Name and location of the project.
    * **ApplicationArea:** The specific application of asphalt (e.g., highway).
    * **SurfaceTechProductsUsed:** The specific Surface Tech products used.
    * **ChallengesFaced:** Any specific challenges or problems addressed.
    * **OutcomesBenefits:** The positive results or benefits achieved by using Surface Tech products.

* **Lifecycle Cost Analysis:**
    * **ProjectScope:** Briefly describe the scope of the cost analysis.
    * **AnalysisPeriod:** The timeframe over which the costs were analyzed.
    * **CostFactorsConsidered:** List the main cost factors included.
    * **KeyFindingsRecommendations:** The main conclusions or recommendations.

* **Technical Product Data:**
    * **ProductName:** The official product name.
    * **ProductCodeSKU:** The unique product identifier.
    * **KeySpecifications:** Important technical specifications (e.g., viscosity grade).
    * **ApplicationInstructions:** Basic instructions on how to use the product.
    * **SDSRef:** Link or reference to the Safety Data Sheet.

**Instructions:**

1. Analyze the provided document text to determine its **DocumentType** from the list above. If the type is unclear, make your best judgment.
2. Extract the **General Metadata** applicable to the identified **DocumentType**.
3. Extract the **Specific Metadata** relevant to the identified **DocumentType**.
4. If a specific piece of metadata cannot be found in the document, indicate it with `<[FieldName]>Not Found</[FieldName]>`.
5. Ensure the output is a well-formed XML document.
6. The root element of the XML should be `<DocumentMetadata>`.
7. Each metadata field should be enclosed within XML tags that exactly match the field names provided in the schema (e.g., `<DocumentTitle>`, `<Authors>`, `<Keywords>`).
8. For fields that can have multiple values (e.g., Authors, Keywords, BATTParametersMeasured, CostFactorsConsidered), enclose each value within its own tag of the same name under the parent tag (e.g., `<Authors><Author>Author 1</Author><Author>Author 2</Author></Authors>`).

**Example Output Format:**

```xml
<DocumentMetadata>
  <DocumentType>Research Paper</DocumentType>
  <DocumentTitle>Impact of Polymer Modification on Asphalt Pavement Performance</DocumentTitle>
  <DatePublishedCreated>2024-03-15</DatePublishedCreated>
  <VersionRevisionNumber>1.0</VersionRevisionNumber>
  <SourceOrigin>Surface Tech Research Lab</SourceOrigin>
  <Keywords>
    <Keyword>Asphalt</Keyword>
    <Keyword>Polymer Modification</Keyword>
    <Keyword>Pavement Performance</Keyword>
  </Keywords>
  <AbstractSummary>This paper investigates the effects of using polymer modifiers...</AbstractSummary>
  <FileFormat>PDF</FileFormat>
  <FileSizeKBMB>2.5 MB</FileSizeKBMB>
  <DocumentIDReferenceNumber>STR-RP-2024-001</DocumentIDReferenceNumber>
  <ConfidentialityLevel>Internal</ConfidentialityLevel>
  <RelevantProductsMaterials>Polymer Modified Asphalt</RelevantProductsMaterials>
  <GeographicRegion>Not Found</GeographicRegion>
  <Authors>
    <Author>John Doe</Author>
    <Author>Jane Smith</Author>
  </Authors>
  <JournalConferenceName>Journal of Asphalt Technology</JournalConferenceName>
  <DOI>10.xxxx/yyyy</DOI>
  <ResearchMethodology>Laboratory testing and field observations</ResearchMethodology>
  <KeyFindings>Polymer modification significantly improved rutting resistance...</KeyFindings>
  <StudyAreaLocation>Not Found</StudyAreaLocation>
</DocumentMetadata>
```
"""

    def extract_metadata(self, text: str, file_path: str, file_format: str) -> Tuple[
        Optional[Dict[str, Any]], Optional[Dict[str, Any]]]:
        """Extract metadata from document text using Gemini API.

        Returns:
            Tuple containing (metadata_dict, raw_response_dict)
        """
        # Add some context about the file to the text
        file_info = f"\nFile Name: {os.path.basename(file_path)}\n"
        file_info += f"File Format: {file_format}\n"
        file_info += f"File Size: {DocumentParser.get_file_size(file_path)}\n\n"

        # Prepare complete text with document content
        complete_text = file_info + text

        # Truncate if too long (Gemini has token limits)
        if len(complete_text) > 100000:  # Arbitrary limit to avoid token issues
            logger.warning(f"Document text is very long ({len(complete_text)} chars). Truncating to 100,000 chars.")
            complete_text = complete_text[:100000] + "\n\n[Content truncated due to length...]"

        try:
            # Call Gemini API with retry logic
            max_retries = 3
            for attempt in range(max_retries):
                try:
                    response = self.model.generate_content(self.prompt + "\n\n" + complete_text)
                    xml_output = response.text

                    # Save the raw API response
                    filename_base = os.path.splitext(os.path.basename(file_path))[0]
                    response_dict = {
                        "prompt": self.prompt,
                        "file_info": file_info,
                        "response": xml_output,
                        "timestamp": time.strftime("%Y-%m-%d %H:%M:%S")
                    }

                    # Create json file with raw response in output directory
                    json_path = os.path.join(self.output_dir, f"{filename_base}.json")
                    with open(json_path, 'w', encoding='utf-8') as f:
                        json.dump(response_dict, f, indent=2, ensure_ascii=False)

                    logger.info(f"Saved raw API response to {json_path}")

                    # Clean up the XML output (remove code blocks if present)
                    if "```xml" in xml_output:
                        xml_output = xml_output.split("```xml")[1].split("```")[0].strip()

                    # Parse the XML to a dictionary
                    metadata_dict = self._xml_to_dict(xml_output)

                    # Add file metadata that we know directly
                    metadata_dict["FileFormat"] = file_format
                    metadata_dict["FileSizeKBMB"] = DocumentParser.get_file_size(file_path)

                    return metadata_dict, response_dict

                except Exception as e:
                    if attempt < max_retries - 1:
                        logger.warning(f"Attempt {attempt + 1} failed: {e}. Retrying...")
                    else:
                        raise

        except Exception as e:
            logger.error(f"Error extracting metadata with Gemini for {file_path}: {e}")
            return None, None

    def _xml_to_dict(self, xml_string: str) -> Dict[str, Any]:
        """Convert XML string to a Python dictionary."""
        try:
            xml_string = xml_string.strip()

            if not xml_string:
                logger.warning("Empty XML string received")
                return {}

            # Wrap the XML in a root element if it doesn't have one
            if not xml_string.startswith("<DocumentMetadata>"):
                xml_string = f"<DocumentMetadata>{xml_string}</DocumentMetadata>"

            root = ET.fromstring(xml_string)
            return self._xml_to_dict_recursive(root)

        except ET.ParseError as e:
            logger.error(f"XML parse error: {e}")
            logger.debug(f"Problematic XML string: {xml_string}")

            # Try to salvage what we can - extract using regex if possible
            import re
            metadata = {}

            # Extract fields using simple pattern matching
            pattern = r"<([^>]+)>(.*?)</\1>"
            matches = re.findall(pattern, xml_string)
            for tag, content in matches:
                if not content.strip().startswith("<"):  # Skip nested elements
                    metadata[tag] = content.strip()

            return metadata

    def _xml_to_dict_recursive(self, element: ET.Element) -> Dict[str, Any]:
        """Recursively convert an XML element to a dictionary."""
        result = {}

        # Process attributes if any
        for key, value in element.attrib.items():
            result[f"@{key}"] = value

        # Process text content
        if element.text and element.text.strip():
            result["#text"] = element.text.strip()

        # Process child elements
        for child in element:
            child_tag = child.tag
            child_dict = self._xml_to_dict_recursive(child)

            if child_tag in result:
                # If this tag already exists, convert to list or append
                if not isinstance(result[child_tag], list):
                    result[child_tag] = [result[child_tag]]
                result[child_tag].append(child_dict)
            else:
                result[child_tag] = child_dict

        # If the element has no children and only text, simplify
        if len(result) == 1 and "#text" in result:
            return result["#text"]

        return result


class ExcelExporter:
    """Class to handle exporting metadata to Excel."""

    def __init__(self, output_filename: str = "metadata_output.xlsx"):
        """Initialize the Excel exporter with the output filename."""
        self.output_filename = output_filename
        self.workbook = openpyxl.Workbook()
        self.sheet = self.workbook.active
        self.sheet.title = "Document Metadata"

        # Define the header row (all possible fields)
        self.header_row = [
            "Filename", "FilePath", "FileFormat", "FileSizeKBMB",
            "DocumentTitle", "DocumentType", "DatePublishedCreated",
            "VersionRevisionNumber", "SourceOrigin", "Keywords",
            "AbstractSummary", "DocumentIDReferenceNumber",
            "ConfidentialityLevel", "RelevantProductsMaterials",
            "GeographicRegion", "Authors", "JournalConferenceName",
            "DOI", "ResearchMethodology", "KeyFindings",
            "StudyAreaLocation", "TestNameType", "TestStandardUsed",
            "DateOfTest", "TestingLaboratory", "EquipmentUsed",
            "MaterialsTested", "TestNameTypeReference", "KeyResultsDataPoints",
            "PassFailCriteria", "ConclusionsInterpretations", "Question",
            "Answer", "RelatedTopics", "DeclaredUnit", "GWP",
            "LCAPractitioner", "PCR", "ValidityPeriod", "ASTMDesignation",
            "YearOfIssue", "TitleOfStandard", "RelevantSections",
            "AramidTypeGrade", "KeyProperties", "Supplier",
            "ApplicationArea", "TestLocation", "BATTParametersMeasured",
            "ResultsSummary", "ProjectNameLocation", "SurfaceTechProductsUsed",
            "ChallengesFaced", "OutcomesBenefits", "ProjectScope",
            "AnalysisPeriod", "CostFactorsConsidered", "KeyFindingsRecommendations",
            "ProductName", "ProductCodeSKU", "KeySpecifications",
            "ApplicationInstructions", "SDSRef", "ProcessingTimeSec"
        ]

        # Add headers to the sheet
        self.sheet.append(self.header_row)

        # Apply formatting to header row
        for cell in self.sheet[1]:
            cell.font = openpyxl.styles.Font(bold=True)
            cell.fill = openpyxl.styles.PatternFill(start_color="D3D3D3", end_color="D3D3D3", fill_type="solid")

    def add_row(self, filename: str, file_path: str, metadata: Dict[str, Any], processing_time: float) -> None:
        """Add a row of metadata to the Excel sheet."""
        row_data = []

        # Add standard file info first
        row_data.append(os.path.basename(filename))  # Filename
        row_data.append(file_path)  # File path

        # Process each field in the header row
        for field in self.header_row[2:]:  # Skip filename and filepath (already added)
            if field == "ProcessingTimeSec":
                row_data.append(f"{processing_time:.2f}")
                continue

            # Extract the field value from metadata
            if field in metadata:
                value = metadata[field]

                # Handle different data types
                if isinstance(value, dict):
                    if "#text" in value:
                        row_data.append(value["#text"])
                    else:
                        # Try to serialize the entire dictionary to JSON
                        try:
                            row_data.append(json.dumps(value, ensure_ascii=False))
                        except Exception as e:
                            logger.warning(f"Could not serialize complex dictionary: {e}")
                            # Fallback to joining parts if serialization fails
                            parts = []
                            for k, v in value.items():
                                if not isinstance(v, dict) and not isinstance(v, list):
                                    parts.append(f"{k}: {v}")
                            row_data.append(", ".join(parts) if parts else str(value))

                elif isinstance(value, list):
                    # Try to serialize the entire list to JSON
                    try:
                        row_data.append(json.dumps(value, ensure_ascii=False))
                    except Exception as e:
                        logger.warning(f"Could not serialize complex list: {e}")
                        # Fallback to joining if serialization fails
                        parts = []
                        for item in value:
                            if isinstance(item, dict) and "#text" in item:
                                parts.append(item["#text"])
                            elif isinstance(item, str):
                                parts.append(item)
                            else:
                                parts.append(str(item))
                        row_data.append(", ".join(parts) if parts else str(value))

                else:
                    # Simple value
                    row_data.append(str(value))
            else:
                row_data.append("Not Found")

        # Add the row to the sheet
        self.sheet.append(row_data)

    def save(self) -> None:
        """Save the Excel workbook."""
        try:
            # Auto-adjust column widths
            for column in self.sheet.columns:
                max_length = 0
                column_letter = openpyxl.utils.get_column_letter(column[0].column)

                for cell in column:
                    if cell.value:
                        cell_length = len(str(cell.value))
                        if cell_length > max_length:
                            max_length = cell_length

                adjusted_width = min(max_length + 2, 50)  # Cap width at 50 to avoid too wide columns
                self.sheet.column_dimensions[column_letter].width = adjusted_width

            # Save the workbook
            self.workbook.save(self.output_filename)
            logger.info(f"Metadata saved to {self.output_filename}")

            return True
        except Exception as e:
            logger.error(f"Error saving Excel file: {e}")
            return False


class MetadataProcessor:
    """Main class to orchestrate the metadata extraction process."""

    def __init__(self, api_key: str, output_dir: str = "output"):
        """Initialize the metadata processor with the API key and output directory."""
        self.document_parser = DocumentParser()
        self.metadata_extractor = MetadataExtractor(api_key, output_dir)
        self.output_dir = output_dir

        # Create output directory if it doesn't exist
        os.makedirs(output_dir, exist_ok=True)

    def process_directory(self, directory_path: str, output_filename: str = "metadata_output.xlsx") -> None:
        """Process all supported files in a directory and export metadata to Excel."""
        directory = Path(directory_path)

        if not directory.exists() or not directory.is_dir():
            logger.error(f"Directory not found or not a directory: {directory_path}")
            raise FileNotFoundError(f"Directory not found: {directory_path}")

        # Get list of supported files
        supported_extensions = [".pdf", ".docx"]
        files = []

        for ext in supported_extensions:
            files.extend(list(directory.glob(f"*{ext}")))

        if not files:
            logger.warning(f"No supported files found in {directory_path}")
            print(f"No supported files found in {directory_path}")
            return

        logger.info(f"Found {len(files)} files to process in {directory_path}")
        print(f"Found {len(files)} files to process")

        # Create Excel exporter
        exporter = ExcelExporter(output_filename)

        # Process each file
        for i, file_path in enumerate(files):
            file_path_str = str(file_path)
            filename = file_path.name

            logger.info(f"Processing file {i + 1}/{len(files)}: {filename}")
            print(f"Processing {i + 1}/{len(files)}: {filename}")

            # Extract text based on file type
            start_time = time.time()
            extracted_text = None
            file_format = file_path.suffix.upper().lstrip('.')

            if file_path.suffix.lower() == ".pdf":
                extracted_text = self.document_parser.extract_text_from_pdf(file_path_str)
            elif file_path.suffix.lower() == ".docx":
                extracted_text = self.document_parser.extract_text_from_docx(file_path_str)

            if not extracted_text:
                logger.warning(f"Could not extract text from {filename}, skipping")
                continue

            # Extract metadata
            metadata, raw_response = self.metadata_extractor.extract_metadata(extracted_text, file_path_str,
                                                                              file_format)
            processing_time = time.time() - start_time

            if metadata:
                # Add to Excel
                exporter.add_row(filename, file_path_str, metadata, processing_time)
                logger.info(f"Successfully processed {filename} in {processing_time:.2f} seconds")
            else:
                logger.error(f"Failed to extract metadata from {filename}")

        # Save Excel file
        if exporter.save():
            print(f"Successfully processed {len(files)} files. Results saved to {output_filename}")
        else:
            print("Error saving results to Excel file")


def main():
    """Main function to parse arguments and run the processor."""
    import time

    parser = argparse.ArgumentParser(description="Extract metadata from documents using Gemini API")
    parser.add_argument("--dir", "-d", required=True, help="Directory containing documents to process")
    parser.add_argument("--output", "-o", default="metadata_output.xlsx", help="Output Excel filename")
    parser.add_argument("--output-dir", default="output",
                        help="Directory to save JSON responses and other output files")
    parser.add_argument("--config", "-c", default="config.json", help="Path to config file with API key")
    args = parser.parse_args()

    # Load API key
    api_key = load_api_key(args.config)
    if not api_key:
        print("Error: No API key found. Please create a config.json file with your API key.")
        print('Example: { "api_key": "YOUR_API_KEY" }')
        return

    try:
        # Create output directory if it doesn't exist
        os.makedirs(args.output_dir, exist_ok=True)

        # Create and run processor
        processor = MetadataProcessor(api_key, args.output_dir)

        # Generate full path for Excel output
        output_path = os.path.join(args.output_dir, args.output)

        processor.process_directory(args.dir, output_path)
    except Exception as e:
        logger.error(f"Error processing documents: {e}")
        print(f"Error: {e}")


if __name__ == "__main__":
    import time

    main()
